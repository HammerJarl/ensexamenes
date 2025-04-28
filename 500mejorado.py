import sqlite3
import docx
from docx import Document
import random
import os
import tkinter as tk
from tkinter import messagebox, filedialog
import logging
from datetime import datetime
import bcrypt

# logging
logging.basicConfig(
    filename='auditoria_examenes.log',
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)

# Contraseña bcrypt
CONTRASEÑA_ADMIN_INICIAL = bcrypt.hashpw("The.Emperor40k".encode('utf-8'), bcrypt.gensalt())
CONTRASEÑA_ADMIN = CONTRASEÑA_ADMIN_INICIAL

# texto es negritas
def es_negrita(run):
    return run.bold or run.font.bold

# bd y tabla
def crear_base_datos():
    conn = sqlite3.connect('base_datos_examenes.db')
    cursor = conn.cursor()
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS preguntas (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            texto_pregunta TEXT NOT NULL,
            opcion_a TEXT NOT NULL,
            opcion_b TEXT NOT NULL,
            opcion_c TEXT NOT NULL,
            opcion_d TEXT NOT NULL,
            opcion_e TEXT NOT NULL,
            respuesta_correcta CHAR(1) NOT NULL
        )
    ''')
    conn.commit()
    conn.close()

# docx y bd
def cargar_preguntas_desde_docx(ruta_docx):
    conn = sqlite3.connect('base_datos_examenes.db')
    cursor = conn.cursor()
    
    doc = Document(ruta_docx)
    pregunta_actual = None
    opciones = []
    etiquetas_opciones = ['a', 'b', 'c', 'd', 'e']
    
    for para in doc.paragraphs:
        if para.text.strip():
            if not pregunta_actual:
                pregunta_actual = para.text.strip()
            else:
                opciones.append(para.text.strip())
                es_correcta = any(es_negrita(run) for run in para.runs)
                if es_correcta:
                    respuesta_correcta = etiquetas_opciones[len(opciones) - 1]
                
                if len(opciones) == 5:
                    cursor.execute('''
                        INSERT INTO preguntas (texto_pregunta, opcion_a, opcion_b, opcion_c, opcion_d, opcion_e, respuesta_correcta)
                        VALUES (?, ?, ?, ?, ?, ?, ?)
                    ''', (pregunta_actual, opciones[0], opciones[1], opciones[2], opciones[3], opciones[4], respuesta_correcta))
                    pregunta_actual = None
                    opciones = []
    
    conn.commit()
    conn.close()
    logging.info(f"Preguntas cargadas desde {ruta_docx}")

def generar_examen(num_preguntas):
    conn = sqlite3.connect('base_datos_examenes.db')
    cursor = conn.cursor()
    
    cursor.execute('SELECT * FROM preguntas')
    todas_preguntas = cursor.fetchall()
    
    if len(todas_preguntas) < num_preguntas:
        return None, f"Solo hay {len(todas_preguntas)} preguntas disponibles, pero se solicitaron {num_preguntas}."
    
    preguntas_seleccionadas = random.sample(todas_preguntas, num_preguntas)
    conn.close()
    return preguntas_seleccionadas, None

# word
def crear_documento_examen(preguntas, nombre_examen, nombre_archivo):
    doc = Document()
    doc.add_heading(f'Examen {nombre_examen}', 0)
    
    for idx, q in enumerate(preguntas, 1):
        opciones = [q[2], q[3], q[4], q[5], q[6]]
        respuesta_correcta = q[7]
        pares_opciones = list(zip(['a', 'b', 'c', 'd', 'e'], opciones))
        random.shuffle(pares_opciones)
        nuevas_etiquetas, nuevas_opciones = zip(*pares_opciones)
        nuevo_indice_correcto = nuevas_etiquetas.index(respuesta_correcta)
        nueva_respuesta_correcta = chr(ord('a') + nuevo_indice_correcto)
        
        doc.add_paragraph(f"{idx}. {q[1]}")
        for etiqueta, opcion in zip(['a', 'b', 'c', 'd', 'e'], nuevas_opciones):
            para = doc.add_paragraph(f"{etiqueta}. {opcion}")
            if etiqueta == nueva_respuesta_correcta:
                for run in para.runs:
                    run.bold = True
    
    doc.save(nombre_archivo)
    logging.info(f"Examen {nombre_examen} generado: {nombre_archivo}")

# limpiar bd
def limpiar_base_datos(contraseña, usuario="Administrador"):
    global CONTRASEÑA_ADMIN
    if not bcrypt.checkpw(contraseña.encode('utf-8'), CONTRASEÑA_ADMIN):
        logging.warning(f"Intento fallido de limpieza de base de datos por {usuario}")
        return False
    
    conn = sqlite3.connect('base_datos_examenes.db')
    cursor = conn.cursor()
    cursor.execute('DELETE FROM preguntas')
    conn.commit()
    conn.close()
    logging.info(f"Base de datos limpiada por {usuario}")
    return True

def cambiar_contraseña(contraseña_actual, nueva_contraseña, usuario="Administrador"):
    global CONTRASEÑA_ADMIN
    if not bcrypt.checkpw(contraseña_actual.encode('utf-8'), CONTRASEÑA_ADMIN):
        logging.warning(f"Intento fallido de cambio de contraseña por {usuario}")
        return False
    
    CONTRASEÑA_ADMIN = bcrypt.hashpw(nueva_contraseña.encode('utf-8'), bcrypt.gensalt())
    logging.info(f"Contraseña de administrador cambiada por {usuario}")
    return True

# Interfaz
class AppExamenes:
    def __init__(self, root):
        self.root = root
        self.root.title("Generador de Exámenes")
        self.root.geometry("600x400")
        
        tk.Label(root, text="Bienvenido al Generador de Exámenes", font=("Arial", 16)).pack(pady=10)
        
        tk.Button(root, text="Cargar Preguntas desde Word", command=self.cargar_preguntas).pack(pady=5)
        tk.Button(root, text="Generar Exámenes", command=self.generar_examenes).pack(pady=5)
        tk.Button(root, text="Limpiar Base de Datos", command=self.limpiar_base).pack(pady=5)
        tk.Button(root, text="Cambiar Contraseña de Administrador", command=self.cambiar_contrasena).pack(pady=5)
        tk.Button(root, text="Salir", command=root.quit).pack(pady=5)
    
    def cargar_preguntas(self):
        ruta_docx = filedialog.askopenfilename(filetypes=[("Documentos Word", "*.docx")])
        if ruta_docx:
            try:
                cargar_preguntas_desde_docx(ruta_docx)
                messagebox.showinfo("Éxito", "Preguntas cargadas exitosamente.")
            except Exception as e:
                messagebox.showerror("Error", f"No se pudieron cargar las preguntas: {e}")
                logging.error(f"Error al cargar preguntas desde {ruta_docx}: {e}")
    
    def generar_examenes(self):
        ventana = tk.Toplevel(self.root)
        ventana.title("Generar Exámenes")
        ventana.geometry("300x200")
        
        tk.Label(ventana, text="Número de preguntas:").pack(pady=10)
        entrada_num = tk.Entry(ventana)
        entrada_num.pack(pady=5)
        
        def generar():
            try:
                num_preguntas = int(entrada_num.get())
                if num_preguntas <= 0:
                    messagebox.showerror("Error", "Ingrese un número positivo.")
                    return
                
                preguntas, error = generar_examen(num_preguntas)
                if error:
                    messagebox.showerror("Error", error)
                    return
                
                random.shuffle(preguntas)
                crear_documento_examen(preguntas, "A", f"Examen_A_{num_preguntas}_preguntas.docx")
                random.shuffle(preguntas)
                crear_documento_examen(preguntas, "B", f"Examen_B_{num_preguntas}_preguntas.docx")
                messagebox.showinfo("Éxito", f"Exámenes generados: Examen_A_{num_preguntas}_preguntas.docx y Examen_B_{num_preguntas}_preguntas.docx")
                ventana.destroy()
            except ValueError:
                messagebox.showerror("Error", "Ingrese un número válido.")
            except Exception as e:
                messagebox.showerror("Error", f"Error al generar exámenes: {e}")
                logging.error(f"Error al generar exámenes: {e}")
        
        tk.Button(ventana, text="Generar", command=generar).pack(pady=10)
    
    def limpiar_base(self):
        ventana = tk.Toplevel(self.root)
        ventana.title("Limpiar Base de Datos")
        ventana.geometry("300x200")
        
        tk.Label(ventana, text="Contraseña de administrador:").pack(pady=10)
        entrada_contraseña = tk.Entry(ventana, show="*")
        entrada_contraseña.pack(pady=5)
        
        def limpiar():
            contraseña = entrada_contraseña.get()
            if limpiar_base_datos(contraseña):
                messagebox.showinfo("Éxito", "Base de datos limpiada exitosamente.")
                ventana.destroy()
            else:
                messagebox.showerror("Error", "Contraseña incorrecta.")
        
        tk.Button(ventana, text="Limpiar", command=limpiar).pack(pady=10)
    
    def cambiar_contrasena(self):
        ventana = tk.Toplevel(self.root)
        ventana.title("Cambiar Contraseña")
        ventana.geometry("300x250")
        
        tk.Label(ventana, text="Contraseña actual:").pack(pady=10)
        entrada_actual = tk.Entry(ventana, show="*")
        entrada_actual.pack(pady=5)
        
        tk.Label(ventana, text="Nueva contraseña:").pack(pady=10)
        entrada_nueva = tk.Entry(ventana, show="*")
        entrada_nueva.pack(pady=5)
        
        def cambiar():
            actual = entrada_actual.get()
            nueva = entrada_nueva.get()
            if not nueva:
                messagebox.showerror("Error", "La nueva contraseña no puede estar vacía.")
                return
            if cambiar_contraseña(actual, nueva):
                messagebox.showinfo("Éxito", "Contraseña cambiada exitosamente.")
                ventana.destroy()
            else:
                messagebox.showerror("Error", "Contraseña actual incorrecta.")
        
        tk.Button(ventana, text="Cambiar", command=cambiar).pack(pady=10)

def main():
    crear_base_datos()
    root = tk.Tk()
    app = AppExamenes(root)
    root.mainloop()

if __name__ == "__main__":
    main()