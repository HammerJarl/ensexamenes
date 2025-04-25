import sqlite3
import docx
from docx import Document
import random
import os
from docx.oxml.ns import qn
from docx.shared import Pt

# texto es negritas
def es_negrita(run):
    return run.bold or run.font.bold

# base de datos y tabla
def crear_base_datos():
    conn = sqlite3.connect('base_datos_examenes.db')
    cursor = conn.cursor()
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS preguntas (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            texto_pregunta TEXT NOT NULL,
            opcion_a TEXT NOT NULL,
            opcion_b TEXT NOT NULL,
            opcion_c TEXT NOT NULL,
            opcion_d TEXT NOT NULL,
            opcion_e TEXT NOT NULL,
            respuesta_correcta CHAR(1) NOT NULL
        )
    ''')
    conn.commit()
    conn.close()

#read docx la resp debe ser siempre a
def cargar_preguntas_docx(ruta_docx):
    conn = sqlite3.connect('base_datos_examenes.db')
    cursor = conn.cursor()

    doc = Document(ruta_docx)
    preguntas = []
    pregunta_actual = None
    opciones = []
    etiquetas_opciones = ['a', 'b', 'c', 'd', 'e']
    
    for para in doc.paragraphs:
        if para.text.strip():
            if not pregunta_actual:
                pregunta_actual = para.text.strip()
            else:
                opciones.append(para.text.strip())
                es_correcta = any(es_negrita(run) for run in para.runs)
                if es_correcta:
                    respuesta_correcta = etiquetas_opciones[len(opciones) - 1]
                
                if len(opciones) == 5:
                    cursor.execute('''
                        INSERT INTO preguntas (texto_pregunta, opcion_a, opcion_b, opcion_c, opcion_d, opcion_e, respuesta_correcta)
                        VALUES (?, ?, ?, ?, ?, ?, ?)
                    ''', (pregunta_actual, opciones[0], opciones[1], opciones[2], opciones[3], opciones[4], respuesta_correcta))
                    pregunta_actual = None
                    opciones = []
    
    conn.commit()
    conn.close()

# Generar un examen con un número de preguntas
def generar_examen(num_preguntas):
    conn = sqlite3.connect('base_datos_examenes.db')
    cursor = conn.cursor()
    
    cursor.execute('SELECT * FROM preguntas')
    todas_preguntas = cursor.fetchall()
    
    if len(todas_preguntas) < num_preguntas:
        print(f"Error: Solo hay {len(todas_preguntas)} preguntas disponibles, pero se solicitaron {num_preguntas}.")
        conn.close()
        return None

    # Randomizador pregs
    preguntas_seleccionadas = random.sample(todas_preguntas, num_preguntas)
    conn.close()
    return preguntas_seleccionadas

# Crear documento Word
def crear_documento_examen(preguntas, nombre_examen, nombre_archivo):
    doc = Document()
    doc.add_heading(f'Examen {nombre_examen}', 0)
    
    for idx, q in enumerate(preguntas, 1):
        # randomizar respuestas
        opciones = [q[2], q[3], q[4], q[5], q[6]]  # a, b, c, d, e
        respuesta_correcta = q[7]
        indice_correcto = ord(respuesta_correcta) - ord('a')
        
        # mezclar respuestas
        pares_opciones = list(zip(['a', 'b', 'c', 'd', 'e'], opciones))
        random.shuffle(pares_opciones)
        nuevas_etiquetas, nuevas_opciones = zip(*pares_opciones)
        
        # Encontrar la nueva respuesta correcta
        nuevo_indice_correcto = nuevas_etiquetas.index(respuesta_correcta)
        nueva_respuesta_correcta = chr(ord('a') + nuevo_indice_correcto)
        
        doc.add_paragraph(f"{idx}. {q[1]}")
        for etiqueta, opcion in zip(['a', 'b', 'c', 'd', 'e'], nuevas_opciones):
            para = doc.add_paragraph(f"{etiqueta}. {opcion}")
            if etiqueta == nueva_respuesta_correcta:
                for run in para.runs:
                    run.bold = False
    
    doc.save(nombre_archivo)
    print(f"Examen guardado como {nombre_archivo}")

# Limpiar BD
def limpiar_base_datos():
    conn = sqlite3.connect('base_datos_examenes.db')
    cursor = conn.cursor()
    cursor.execute('DELETE FROM preguntas')
    conn.commit()
    conn.close()
    print("Base de datos limpiada.")

# Menú principal
def main():
    crear_base_datos()
    
    while True:
        print("\n=== Menú Generador de Exámenes ===")
        print("1. Cargar preguntas desde un documento Word")
        print("2. Generar exámenes")
        print("3. Limpiar base de datos (Administrador)")
        print("4. Salir")
        
        opcion = input("Ingrese su opción (1-4): ")
        
        if opcion == '1':
            ruta_docx = input("Ingrese la ruta del documento Word: ")
            if os.path.exists(ruta_docx):
                cargar_preguntas_docx(ruta_docx)  
                print("Preguntas cargadas exitosamente.")
            else:
                print("Archivo no encontrado.")
        
        elif opcion == '2':
            try:
                num_preguntas = int(input("Ingrese el número de preguntas para el examen (ej. 50, 100, 200): "))
                if num_preguntas <= 0:
                    print("Por favor, ingrese un número positivo.")
                    continue
                
                preguntas = generar_examen(num_preguntas)
                if preguntas:
                    random.shuffle(preguntas)  # Mezclar para Examen A
                    crear_documento_examen(preguntas, "A", f"Examen_A_{num_preguntas}_preguntas.docx")
                    
                    random.shuffle(preguntas)  # Mezclar para Examen B
                    crear_documento_examen(preguntas, "B", f"Examen_B_{num_preguntas}_preguntas.docx")
            except ValueError:
                print("Por favor, ingrese un número válido.")
        
        elif opcion == '3':
            contraseña_admin = input("Ingrese la contraseña de administrador: ")
            if contraseña_admin == "The.Emperor40k":
                limpiar_base_datos()
            else:
                print("Contraseña incorrecta.")
        
        elif opcion == '4':
            print("Saliendo del programa.")
            break
        
        else:
            print("Opción inválida. Intente de nuevo.")

if __name__ == "__main__":
    main()