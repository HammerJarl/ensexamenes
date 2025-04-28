import sqlite3
import docx
from docx import Document
import random
import os
import tkinter as tk
from tkinter import messagebox, filedialog, ttk
import logging
from datetime import datetime
import bcrypt
from PIL import Image, ImageTk
import sys
import time

logging.basicConfig(
    filename='auditoria_examenes.log',
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)

CONTRASEÑA_ADMIN_INICIAL = bcrypt.hashpw("The.Emperor40k".encode('utf-8'), bcrypt.gensalt())
CONTRASEÑA_ADMIN = CONTRASEÑA_ADMIN_INICIAL


def resource_path(relative_path):
    """Obtiene la ruta absoluta del recurso, funciona tanto para desarrollo como para el .exe empaquetado."""
    if hasattr(sys, '_MEIPASS'):
        return os.path.join(sys._MEIPASS, relative_path)
    else:
        return os.path.join(os.path.abspath("."), relative_path)

def convertir_a_ico(png_path, ico_path):
    try:
        img = Image.open(png_path)
        img = img.resize((512, 512), Image.LANCZOS)
        img.save(ico_path, format='ICO')
        logging.info(f"Imagen convertida de {png_path} a {ico_path}")
        return True
    except Exception as e:
        logging.error(f"Error al convertir {png_path} a ICO: {e}")
        return False

def preparar_iconos():
    scripts_dir = "imagenes"
    icono_png = os.path.join(scripts_dir, "ironpriest.png")
    icono_ico = "ironpriest.ico"

    if os.path.exists(icono_png):
        if not os.path.exists(icono_ico):
            if convertir_a_ico(icono_png, icono_ico):
                logging.info(f"Icono generado: {icono_ico}")
            else:
                logging.warning(f"No se pudo generar {icono_ico}. Se usará el icono predeterminado.")
        else:
            logging.info(f"Icono ya existe: {icono_ico}")
    else:
        logging.warning(f"{icono_png} no encontrado. Se usará el icono predeterminado.")

def es_negrita(run):
    return run.bold or run.font.bold

def es_cursiva(run):
    return run.italic or run.font.italic

def crear_base_datos():
    conn = sqlite3.connect('base_datos_examenes.db')
    cursor = conn.cursor()
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS preguntas (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            texto_pregunta TEXT NOT NULL,
            opcion_a TEXT NOT NULL,
            opcion_b TEXT NOT NULL,
            opcion_c TEXT NOT NULL,
            opcion_d TEXT NOT NULL,
            opcion_e TEXT NOT NULL,
            respuesta_correcta CHAR(1) NOT NULL
        )
    ''')
    conn.commit()
    conn.close()

def cargar_preguntas_desde_docx(ruta_docx):
    conn = sqlite3.connect('base_datos_examenes.db')
    cursor = conn.cursor()
    
    doc = Document(ruta_docx)
    pregunta_actual = None
    opciones = []
    respuesta_correcta = None
    etiquetas_opciones = ['a', 'b', 'c', 'd', 'e']
    
    for para in doc.paragraphs:
        texto = para.text.strip()
        if not texto:
            continue
        
        print(f"Procesando párrafo: {texto}")
        
        if (texto.startswith('¿') and texto.endswith('?') and 
            any(es_cursiva(run) for run in para.runs if run.italic is not None)):
            if pregunta_actual and len(opciones) == 5:
                try:
                    cursor.execute('''
                        INSERT INTO preguntas (texto_pregunta, opcion_a, opcion_b, opcion_c, opcion_d, opcion_e, respuesta_correcta)
                        VALUES (?, ?, ?, ?, ?, ?, ?)
                    ''', (pregunta_actual, opciones[0], opciones[1], opciones[2], opciones[3], opciones[4], respuesta_correcta))
                    logging.info(f"Pregunta '{pregunta_actual}' cargada correctamente.")
                except Exception as e:
                    logging.error(f"Error al insertar pregunta '{pregunta_actual}': {e}")
            
            pregunta_actual = texto
            opciones = []
            respuesta_correcta = None
            print(f"Nueva pregunta: {pregunta_actual}")
            continue
        
        if (texto.startswith('*') and texto.endswith('*') and 
            any(es_negrita(run) for run in para.runs if run.bold is not None)):
            if opciones:
                logging.error(f"Pregunta '{pregunta_actual}' ya tiene opciones antes de la respuesta correcta. Saltando.")
                pregunta_actual = None
                opciones = []
                respuesta_correcta = None
                continue
            opciones.append(texto[1:-1].strip())
            respuesta_correcta = 'a'
            print(f"Respuesta correcta: {texto[1:-1].strip()}")
            continue
        
        if pregunta_actual and respuesta_correcta:
            opciones.append(texto)
            print(f"Opción incorrecta: {texto}")
        
        if len(opciones) == 5:
            if not respuesta_correcta:
                logging.error(f"Pregunta '{pregunta_actual}' no tiene respuesta correcta. Saltando.")
                pregunta_actual = None
                opciones = []
                respuesta_correcta = None
                continue
            
            try:
                cursor.execute('''
                    INSERT INTO preguntas (texto_pregunta, opcion_a, opcion_b, opcion_c, opcion_d, opcion_e, respuesta_correcta)
                    VALUES (?, ?, ?, ?, ?, ?, ?)
                ''', (pregunta_actual, opciones[0], opciones[1], opciones[2], opciones[3], opciones[4], respuesta_correcta))
                logging.info(f"Pregunta '{pregunta_actual}' cargada correctamente.")
            except Exception as e:
                logging.error(f"Error al insertar pregunta '{pregunta_actual}': {e}")
            
            pregunta_actual = None
            opciones = []
            respuesta_correcta = None
    
    if pregunta_actual and len(opciones) == 5 and respuesta_correcta:
        try:
            cursor.execute('''
                INSERT INTO preguntas (texto_pregunta, opcion_a, opcion_b, opcion_c, opcion_d, opcion_e, respuesta_correcta)
                VALUES (?, ?, ?, ?, ?, ?, ?)
            ''', (pregunta_actual, opciones[0], opciones[1], opciones[2], opciones[3], opciones[4], respuesta_correcta))
            logging.info(f"Pregunta '{pregunta_actual}' cargada correctamente.")
        except Exception as e:
            logging.error(f"Error al insertar pregunta '{pregunta_actual}': {e}")
    
    if pregunta_actual or opciones:
        logging.warning(f"Documento incompleto: pregunta '{pregunta_actual}' con {len(opciones)} opciones no procesada.")
    
    conn.commit()
    conn.close()
    logging.info(f"Proceso de carga completado para {ruta_docx}")

def generar_examen(num_preguntas):
    conn = sqlite3.connect('base_datos_examenes.db')
    cursor = conn.cursor()
    
    cursor.execute('SELECT * FROM preguntas')
    todas_preguntas = cursor.fetchall()
    
    if len(todas_preguntas) < num_preguntas:
        return None, f"Solo hay {len(todas_preguntas)} preguntas disponibles, pero se solicitaron {num_preguntas}."
    
    preguntas_seleccionadas = random.sample(todas_preguntas, num_preguntas)
    conn.close()
    return preguntas_seleccionadas, None

def crear_documento_examen(preguntas, nombre_examen, nombre_archivo):
    doc = Document()
    doc.add_heading(f'Examen {nombre_examen}', 0)
    
    for idx, q in enumerate(preguntas, 1):
        opciones = [q[2], q[3], q[4], q[5], q[6]]
        respuesta_correcta = q[7]
        
        pares_opciones = list(zip(['a', 'b', 'c', 'd', 'e'], opciones))
        random.shuffle(pares_opciones)
        nuevas_etiquetas, nuevas_opciones = zip(*pares_opciones)
        
        nuevo_indice_correcto = nuevas_etiquetas.index(respuesta_correcta)
        nueva_respuesta_correcta = chr(ord('a') + nuevo_indice_correcto)
        
        para_pregunta = doc.add_paragraph(f"{idx}. {q[1]}")
        for run in para_pregunta.runs:
            run.italic = True
        
        for etiqueta, opcion in zip(['a', 'b', 'c', 'd', 'e'], nuevas_opciones):
            para = doc.add_paragraph(f"{etiqueta}. {opcion}")
            if etiqueta == nueva_respuesta_correcta:
                for run in para.runs:
                    run.bold = True
    
    doc.save(nombre_archivo)
    logging.info(f"Examen {nombre_examen} generado: {nombre_archivo}")

def limpiar_base_datos(contraseña, usuario="Administrador"):
    global CONTRASEÑA_ADMIN
    if not bcrypt.checkpw(contraseña.encode('utf-8'), CONTRASEÑA_ADMIN):
        logging.warning(f"Intento fallido de limpieza de base de datos por {usuario}")
        return False
    
    conn = sqlite3.connect('base_datos_examenes.db')
    cursor = conn.cursor()
    cursor.execute('DELETE FROM preguntas')
    conn.commit()
    conn.close()
    logging.info(f"Base de datos limpiada por {usuario}")
    return True

def cambiar_contraseña(contraseña_actual, nueva_contraseña, usuario="Administrador"):
    global CONTRASEÑA_ADMIN
    if not bcrypt.checkpw(contraseña_actual.encode('utf-8'), CONTRASEÑA_ADMIN):
        logging.warning(f"Intento fallido de cambio de contraseña por {usuario}")
        return False
    
    CONTRASEÑA_ADMIN = bcrypt.hashpw(nueva_contraseña.encode('utf-8'), bcrypt.gensalt())
    logging.info(f"Contraseña de administrador cambiada por {usuario}")
    return True

def mostrar_pantalla_carga(root):
    splash = tk.Toplevel(root)
    splash.overrideredirect(True)

    splash_width = 440
    splash_height = 650
    splash.geometry(f"{splash_width}x{splash_height}")

    screen_width = root.winfo_screenwidth()
    screen_height = root.winfo_screenheight()
    x_position = (screen_width - splash_width) // 2
    y_position = (screen_height - splash_height) // 2
    splash.geometry(f"{splash_width}x{splash_height}+{x_position}+{y_position}")

    splash.configure(bg="#FF00FF")
    try:
        splash.attributes('-transparentcolor', '#FF00FF')
    except Exception as e:
        logging.warning(f"Transparencia no soportada en esta plataforma: {e}")
        splash.configure(bg="#aac7e3")

    splash_image_path = resource_path(os.path.join("imagenes", "IPlogo.png"))
    try:
        splash_img = Image.open(splash_image_path)
        splash_img = splash_img.resize((440, 650), Image.LANCZOS)
        splash_photo = ImageTk.PhotoImage(splash_img)
        splash_label = tk.Label(splash, image=splash_photo, bg="#FF00FF")
        splash_label.image = splash_photo
        splash_label.pack(expand=True, fill="both")
    except Exception as e:
        logging.error(f"Error al cargar la imagen de la pantalla de carga ({splash_image_path}): {e}")
        splash_label = tk.Label(splash, text="Cargando...", font=("Helvetica", 20), bg="#FF00FF", fg="#003087")
        splash_label.pack(expand=True, fill="both")

    progress = ttk.Progressbar(splash, orient="horizontal", length=300, mode="determinate")
    progress.place(relx=0.5, rely=0.9, anchor="center")

    steps = 100
    delay = 6000 // steps
    for i in range(steps + 1):
        progress['value'] = i
        splash.update()
        time.sleep(delay / 1000)

    splash.destroy()
    
class AppExamenes:
    def __init__(self, root):
        self.root = root
        self.root.title("Exámenes ENS")
        self.root.geometry("800x600")
        self.root.configure(bg="#aac7e3")
        
        icono_ventana_path = resource_path("ironpriest.ico")
        if os.path.exists(icono_ventana_path):
            try:
                self.root.iconbitmap(icono_ventana_path)
                logging.info(f"Icono de la ventana cargado: {icono_ventana_path}")
            except Exception as e:
                logging.error(f"Error al cargar el icono de la ventana ({icono_ventana_path}): {e}")
        else:
            logging.warning(f"Icono de la ventana no encontrado: {icono_ventana_path}. Usando icono predeterminado.")
        
        self.style = ttk.Style()
        self.style.configure("TButton", font=("Helvetica", 12), padding=10)
        self.style.map("TButton",
                       background=[('active', '#d1e7ff'), ('!active', '#e6f0fa')],
                       foreground=[('active', '#003087'), ('!active', '#003087')])
        
        frame = ttk.Frame(root, padding="20", style="Custom.TFrame")
        frame.pack(expand=True, fill="both")
        
        self.style.configure("Custom.TFrame", background="#f0f4f8")
        
        ttk.Label(frame, text="Generador de Exámenes", font=("Helvetica", 20, "bold"), background="#f0f4f8", foreground="#003087").pack(pady=20)
        
        ttk.Button(frame, text="Cargar Preguntas", command=self.cargar_preguntas).pack(pady=10, fill="x")
        ttk.Button(frame, text="Generar Exámenes", command=self.generar_examenes).pack(pady=10, fill="x")
        ttk.Button(frame, text="Limpiar Base de Datos", command=self.limpiar_base).pack(pady=10, fill="x")
        ttk.Button(frame, text="Cambiar Contraseña de Administrador", command=self.cambiar_contrasena).pack(pady=10, fill="x")
        ttk.Button(frame, text="Salir", command=root.quit).pack(pady=10, fill="x")
        #der
        ens_logo_path = resource_path(os.path.join("imagenes", "ens_logo.png"))
        try:
            ens_img = Image.open(ens_logo_path)
            ens_img = ens_img.resize((90, 120), Image.LANCZOS)
            ens_photo = ImageTk.PhotoImage(ens_img)
            ens_label = tk.Label(frame, image=ens_photo, bg="#f0f4f8")
            ens_label.image = ens_photo
            ens_label.place(relx=1.0, rely=1.0, anchor="se")
        except Exception as e:
            logging.error(f"Error al cargar el logo de la ENS ({ens_logo_path}): {e}")
            ens_label = tk.Label(frame, text="Logo no disponible", font=("Helvetica", 10), bg="#f0f4f8", fg="#003087")
            ens_label.place(relx=1.0, rely=1.0, anchor="se")
        #izq
        it_logo_path = resource_path(os.path.join("imagenes", "ironpriest.png"))
        try:
            it_img = Image.open(it_logo_path)
            it_img = it_img.resize((110, 120), Image.LANCZOS)
            it_photo = ImageTk.PhotoImage(it_img)
            it_label = tk.Label(frame, image=it_photo, bg="#f0f4f8")
            it_label.image = it_photo
            it_label.place(relx=0.0, rely=1.0, anchor="sw")
        except Exception as e:
            logging.error(f"Error al cargar el logo de la ENS ({it_logo_path}): {e}")
            it_label = tk.Label(frame, text="Logo no disponible", font=("Helvetica", 10), bg="#f0f4f8", fg="#003087")
            it_label.place(relx=0.0, rely=1.0, anchor="sw")
        #mid
        cr_logo_path = resource_path(os.path.join("imagenes", "copyright.png"))
        try:
            cr_img = Image.open(cr_logo_path)
            cr_img = cr_img.resize((250, 125), Image.LANCZOS)
            cr_photo = ImageTk.PhotoImage(cr_img)
            cr_label = tk.Label(frame, image=cr_photo, bg="#f0f4f8")
            cr_label.image = cr_photo
            cr_label.place(relx=0.5, rely=1.0, anchor="s")
        except Exception as e:
            logging.error(f"Error al cargar el logo de la ENS ({cr_logo_path}): {e}")
            cr_label = tk.Label(frame, text="Logo no disponible", font=("Helvetica", 10), bg="#f0f4f8", fg="#003087")
            cr_label.place(relx=0.5, rely=1.0, anchor="sw")

    def cargar_preguntas(self):
        ruta_docx = filedialog.askopenfilename(filetypes=[("Documentos Word", "*.docx")])
        if ruta_docx:
            try:
                cargar_preguntas_desde_docx(ruta_docx)
                messagebox.showinfo("Éxito", "Preguntas cargadas exitosamente.")
            except Exception as e:
                messagebox.showerror("Error", f"No se pudieron cargar las preguntas: {e}")
                logging.error(f"Error al cargar preguntas desde {ruta_docx}: {e}")
    
    def generar_examenes(self):
        ventana = tk.Toplevel(self.root)
        ventana.title("Generar Exámenes")
        ventana.geometry("400x300")
        ventana.configure(bg="#f0f4f8")
        
        icono_ventana_path = resource_path("ironpriest.ico")
        if os.path.exists(icono_ventana_path):
            try:
                ventana.iconbitmap(icono_ventana_path)
                logging.info(f"Icono de la ventana secundaria cargado: {icono_ventana_path}")
            except Exception as e:
                logging.error(f"Error al cargar el icono de la ventana secundaria ({icono_ventana_path}): {e}")
        else:
            logging.warning(f"Icono de la ventana secundaria no encontrado: {icono_ventana_path}")
        
        frame = ttk.Frame(ventana, padding="20", style="Custom.TFrame")
        frame.pack(expand=True, fill="both")
        
        ttk.Label(frame, text="Número de preguntas:", font=("Helvetica", 12), background="#f0f4f8").pack(pady=10)
        entrada_num = ttk.Entry(frame, font=("Helvetica", 12))
        entrada_num.pack(pady=5, fill="x")
        
        def generar():
            try:
                num_preguntas = int(entrada_num.get())
                if num_preguntas <= 0:
                    messagebox.showerror("Error", "Ingrese un número positivo.")
                    return
                
                preguntas, error = generar_examen(num_preguntas)
                if error:
                    messagebox.showerror("Error", error)
                    return
                
                random.shuffle(preguntas)
                nombre_archivo_a = f"Examen_A_{num_preguntas}_preguntas.docx"
                crear_documento_examen(preguntas, "A", nombre_archivo_a)
                
                random.shuffle(preguntas)
                nombre_archivo_b = f"Examen_B_{num_preguntas}_preguntas.docx"
                crear_documento_examen(preguntas, "B", nombre_archivo_b)
                
                messagebox.showinfo("Éxito", f"Exámenes generados: {nombre_archivo_a} y {nombre_archivo_b}")
                ventana.destroy()
            except ValueError:
                messagebox.showerror("Error", "Ingrese un número válido.")
            except Exception as e:
                messagebox.showerror("Error", f"Error al generar exámenes: {e}")
                logging.error(f"Error al generar exámenes: {e}")
        
        ttk.Button(frame, text="Generar", command=generar).pack(pady=10, fill="x")
    
    def limpiar_base(self):
        ventana = tk.Toplevel(self.root)
        ventana.title("Limpiar Base de Datos")
        ventana.geometry("400x300")
        ventana.configure(bg="#f0f4f8")
        
        icono_ventana_path = resource_path("ironpriest.ico")
        if os.path.exists(icono_ventana_path):
            try:
                ventana.iconbitmap(icono_ventana_path)
                logging.info(f"Icono de la ventana secundaria cargado: {icono_ventana_path}")
            except Exception as e:
                logging.error(f"Error al cargar el icono de la ventana secundaria ({icono_ventana_path}): {e}")
        else:
            logging.warning(f"Icono de la ventana secundaria no encontrado: {icono_ventana_path}")
        
        frame = ttk.Frame(ventana, padding="20", style="Custom.TFrame")
        frame.pack(expand=True, fill="both")
        
        ttk.Label(frame, text="Contraseña de administrador:", font=("Helvetica", 12), background="#f0f4f8").pack(pady=10)
        entrada_contraseña = ttk.Entry(frame, show="*", font=("Helvetica", 12))
        entrada_contraseña.pack(pady=5, fill="x")
        
        def limpiar():
            contraseña = entrada_contraseña.get()
            if limpiar_base_datos(contraseña):
                messagebox.showinfo("Éxito", "Base de datos limpiada exitosamente.")
                ventana.destroy()
            else:
                messagebox.showerror("Error", "Contraseña incorrecta.")
        
        ttk.Button(frame, text="Limpiar", command=limpiar).pack(pady=10, fill="x")
    
    def cambiar_contrasena(self):
        ventana = tk.Toplevel(self.root)
        ventana.title("Cambiar Contraseña")
        ventana.geometry("400x350")
        ventana.configure(bg="#f0f4f8")
        
        icono_ventana_path = resource_path("ironpriest.ico")
        if os.path.exists(icono_ventana_path):
            try:
                ventana.iconbitmap(icono_ventana_path)
                logging.info(f"Icono de la ventana secundaria cargado: {icono_ventana_path}")
            except Exception as e:
                logging.error(f"Error al cargar el icono de la ventana secundaria ({icono_ventana_path}): {e}")
        else:
            logging.warning(f"Icono de la ventana secundaria no encontrado: {icono_ventana_path}")
        
        frame = ttk.Frame(ventana, padding="20", style="Custom.TFrame")
        frame.pack(expand=True, fill="both")
        
        ttk.Label(frame, text="Contraseña actual:", font=("Helvetica", 12), background="#f0f4f8").pack(pady=10)
        entrada_actual = ttk.Entry(frame, show="*", font=("Helvetica", 12))
        entrada_actual.pack(pady=5, fill="x")
        
        ttk.Label(frame, text="Nueva contraseña:", font=("Helvetica", 12), background="#f0f4f8").pack(pady=10)
        entrada_nueva = ttk.Entry(frame, show="*", font=("Helvetica", 12))
        entrada_nueva.pack(pady=5, fill="x")
        
        def cambiar():
            actual = entrada_actual.get()
            nueva = entrada_nueva.get()
            if not nueva:
                messagebox.showerror("Error", "La nueva contraseña no puede estar vacía.")
                return
            if cambiar_contraseña(actual, nueva):
                messagebox.showinfo("Éxito", "Contraseña cambiada exitosamente.")
                ventana.destroy()
            else:
                messagebox.showerror("Error", "Contraseña actual incorrecta.")
        
        ttk.Button(frame, text="Cambiar", command=cambiar).pack(pady=10, fill="x")

def main():
    preparar_iconos()

    root = tk.Tk()
    root.withdraw()

    mostrar_pantalla_carga(root)

    root.deiconify()
    crear_base_datos()
    app = AppExamenes(root)
    root.mainloop()

if __name__ == "__main__":
    main()